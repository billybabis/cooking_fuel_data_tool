"""
Generate USER_GUIDE.docx.

*** WARNING — READ BEFORE RUNNING ***
The USER_GUIDE.docx in this repo has real screenshots pasted into it by hand, in place
of the "[ Screenshot placeholder ]" paragraphs this script emits. Running this script
OVERWRITES that file and DESTROYS those screenshots — they are not stored anywhere else
in the repo.

Prefer editing USER_GUIDE.docx directly (python-docx can modify it in place without
touching the images). Only run this script to rebuild the guide from scratch, and
re-insert the screenshots afterwards.

The text below is kept in sync with the shipped .docx so the two do not drift.

Run: python generate_user_guide_docx.py
Output: USER_GUIDE.docx (overwrites if it exists)

Requires python-docx (already installed in this project's env).
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---------- helpers ----------

def add_para(doc, text, *, italic=False, bold=False, size=None, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    if align is not None:
        p.alignment = align
    return p


def add_screenshot_placeholder(doc, description, filename):
    """A styled paragraph that flags where a screenshot should go."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run_label = p.add_run("[ Screenshot placeholder ] ")
    run_label.bold = True
    run_label.font.color.rgb = RGBColor(0x99, 0x66, 0x00)
    run_desc = p.add_run(description)
    run_desc.italic = True
    run_path = p.add_run(f"  ({filename})")
    run_path.italic = True
    run_path.font.size = Pt(9)
    run_path.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def add_table(doc, headers, rows):
    """Build a bordered table. headers: list[str]; rows: list[list[str]]."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.bold = True
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            t.rows[r].cells[c].text = str(val)
    # small bottom spacing
    doc.add_paragraph()


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)


# ---------- build the document ----------

doc = Document()

# Base style: a touch larger and a sensible font
for s in doc.styles:
    if s.name == "Normal":
        s.font.name = "Calibri"
        s.font.size = Pt(11)

# Title
title = doc.add_heading("Cooking Fuel Data Tool — User Guide", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT

doc.add_paragraph(
    "This guide walks through every screen and control in the Cooking Fuel Data Tool. "
    "It is intended for analysts and modellers who will use the tool."
)
doc.add_paragraph(
    "The tool always opens to Outputs with the default selections applied. The two "
    "top-level views (Outputs / Inputs) sit as folder tabs at the top of the body; "
    "the sidebar on the left holds the filters and applies to both views."
)
add_screenshot_placeholder(
    doc,
    'Full window on first load, with the sidebar visible and the "Total fuel consumption" sub-tab open',
    "docs/screenshots/00-overview.png",
)

# ---------- Section 1 ----------

doc.add_heading("Section 1 — Default Output: Total Fuel Consumption", level=1)
doc.add_paragraph(
    "On launch you see the Total fuel consumption sub-tab inside the Outputs view. "
    "The big table is the headline output of the tool: estimated fuel users and total "
    "consumption for every selected (country, area, fuel, year) combination."
)
add_screenshot_placeholder(
    doc,
    "Sidebar + consumption table + download buttons",
    "docs/screenshots/01-default-consumption-tab.png",
)

doc.add_heading("1.1 Sidebar filters", level=2)
doc.add_paragraph(
    "The left-side panel controls what data the rest of the app shows. Changes apply "
    "immediately on rerun — no \"apply\" button needed."
)
add_table(
    doc,
    ["Control", "What it does", "Notes"],
    [
        ["Countries",
         "Multi-select of all countries present in the fuel-share dataset",
         "Required — the body shows a prompt to pick at least one country before any data is rendered."],
        ["Fuel types",
         "Multi-select of fuel categories: fuelwood, charcoal, coal, gas, kerosene, electric, biogas, ethanol, pellets, imp_fuelwood, imp_charcoal",
         "Defaults to the six fuels covered by the UN/WHO household survey (fuelwood, charcoal, coal, gas, kerosene, electric). The other five — biogas, ethanol, pellets, imp_fuelwood, imp_charcoal — are not in that survey, so they carry a 0% population share everywhere and start deselected. Select one here and enter shares in the Fuel shares tab (§ 2.1) to bring it into the output."],
        ["Areas",
         "Multi-select: urban, rural, overall",
         "overall is derived as urban + rural (see § 1.3). Default is all three."],
        ["End year",
         "Slider snapping to one of 2020 / 2030 / 2035 / 2040 / 2050",
         "Start year is fixed at 2000. Default end year is 2050."],
    ],
)
add_screenshot_placeholder(
    doc,
    "Close-up of the sidebar with all four filter widgets",
    "docs/screenshots/02-sidebar-filters.png",
)

doc.add_heading("1.2 The output table", level=2)
doc.add_paragraph(
    "The table is a single row per (iso3, country, area, fuel, year) combination. Columns:"
)
add_table(
    doc,
    ["Column", "Meaning", "Units"],
    [
        ["iso3, country, area, fuel, year",
         "Identifiers — area is one of urban/rural/overall; fuel is lowercase", "—"],
        ["num_fuel_users_thousands",
         "Number of people using this fuel in the given (country, area, year)",
         "Thousands of people (UN WPP convention)"],
        ["per_capita_fuel_cons",
         "Per-capita consumption rate from the per-capita dataset",
         "Varies by fuel: MWh/person-year for electric; oven-dry tons/person-year for fuelwood/imp_fuelwood; tons/person-year for others"],
        ["fuel_cons_tons",
         "Total consumption: num_fuel_users_thousands × per_capita_fuel_cons × 1,000. "
         "Charcoal rows are additionally multiplied by the kiln yield (§ 3.2), and wood/charcoal "
         "rows by the non-residential uplift if it is enabled (§ 3.3) — so for those fuels the "
         "plain three-term identity does not hold.",
         "Absolute tons (rounded to whole numbers in the display and exports)"],
    ],
)
doc.add_paragraph(
    "The region column exists in the backend (used to join per-capita rates) but is hidden "
    "from the display so the table stays readable."
)
add_screenshot_placeholder(
    doc,
    "Consumption table with caption visible below",
    "docs/screenshots/03-consumption-table.png",
)

doc.add_heading("Units quick-reference", level=3)
doc.add_paragraph(
    "num_fuel_users_thousands is in thousands of people, but fuel_cons_tons is in "
    "absolute tons. The calculation applies a ×1,000 rescale internally so the consumption "
    "value lands in absolute tons (more useful for emissions and policy). Because of that, "
    "the identity users × per_capita = fuel_cons_tons does NOT hold without the additional "
    "×1,000 — see the inline caption beneath the table and the metadata sheet of the Excel "
    "export for the explicit note."
)

doc.add_heading("1.3 How overall is computed", level=2)
doc.add_paragraph(
    "The overall row for each (country, fuel, year) is computed as the sum of the urban "
    "and rural rows for that combination — not taken from the source's \"Overall\" series. "
    "This guarantees overall.num_fuel_users_thousands = urban.num_fuel_users_thousands + "
    "rural.num_fuel_users_thousands and the same for fuel_cons_tons. The source's Overall "
    "rows from the WHO/Stoner dataset are intentionally discarded; the modelled \"Overall\" "
    "share doesn't always equal the population-weighted average of the urban and rural "
    "shares, and the tool prefers the arithmetic identity for consistency."
)

doc.add_heading("1.4 Downloading the data", level=2)
doc.add_paragraph(
    "Below the table are two download buttons side-by-side:"
)
add_screenshot_placeholder(
    doc,
    "Download buttons + filename input",
    "docs/screenshots/04-download-buttons.png",
)
add_table(
    doc,
    ["Button", "Format", "Contents"],
    [
        ["Download Excel Workbook", ".xlsx",
         "Two sheets: Metadata & Sources (run parameters, citations, units notes) and Fuel Consumption Data (the table)."],
        ["Download CSV", ".csv",
         "Single file. The first two rows are a parameter preamble — row 1 variable names, row 2 the matching values — carrying the run parameters and total demand per fuel (see \"The CSV parameter preamble\" below). Row 3 onward is the data table (header + rows). Designed for ingestion by downstream tools that expect a named-parameter header above the data."],
    ],
)
doc.add_paragraph(
    "The text input above the buttons sets the filename stem — both downloads use it (with "
    ".xlsx and .csv extensions respectively). Illegal filename characters are stripped automatically."
)

doc.add_heading("The CSV parameter preamble", level=3)
doc.add_paragraph(
    "The CSV export opens with a two-row parameter preamble: row 1 holds variable names, "
    "row 2 holds the matching values, column for column. The consumption data table begins "
    "on row 3. This lets a downstream tool read the run's settings and headline results "
    "straight off the top of the file without re-aggregating the data rows."
)
add_code(
    doc,
    "end_year,efchratio,nonres_wood_pct,nonres_charcoal_pct,demand_yr_start,demand_yr_end,"
    "demand_fw,demand_ch,demand_coal,demand_gas,demand_kero,demand_elec,demand_biogas,"
    "demand_eth,demand_pel,demand_impfw,demand_impch",
)
add_code(doc, "2050,6,10,20,2026,2035,158876572,7045763,0,8530677,942834,148107,0,0,0,0,0")
doc.add_paragraph("The parameter columns:")
add_table(
    doc,
    ["Column", "Meaning", "Units"],
    [
        ["end_year", "Last projection year, from the sidebar slider.", "year"],
        ["efchratio",
         "Wood-to-charcoal kiln yield (§ 3.2), already applied to charcoal rows.",
         "kg wood / kg charcoal"],
        ["nonres_wood_pct",
         "Non-residential uplift applied to fuelwood and imp_fuelwood. 0 when the option is "
         "switched off — a 0% uplift is equivalent to not applying it, so there is no separate "
         "on/off flag.",
         "percent"],
        ["nonres_charcoal_pct",
         "Non-residential uplift applied to charcoal and imp_charcoal. 0 when switched off.",
         "percent"],
        ["demand_yr_start, demand_yr_end",
         "The year window the demand_* totals were summed over. Normally 2026–2035; clipped if "
         "the end-year filter stops earlier.",
         "year"],
        ["demand_<fuel>",
         "Total consumption of that fuel over the demand window — one value per fuel. See the "
         "mapping table below.",
         "tons (MWh for demand_elec)"],
    ],
)

doc.add_heading("Demand columns", level=4)
doc.add_paragraph(
    "Each demand_<fuel> value is the total consumption of that fuel summed over the demand "
    "window, across the countries currently selected in the sidebar, using the overall area "
    "rows (overall = urban + rural, so this is a national total and is not double-counted "
    "against the urban and rural rows). Values reflect everything applied to the data rows "
    "below: inline edits, the kiln yield, and the non-residential uplift."
)
doc.add_paragraph(
    "All eleven fuels are always written, in the fixed order below, with 0 for any fuel that "
    "is deselected in the sidebar or carries a zero share. The preamble layout is therefore "
    "identical between exports, so a downstream parser can rely on a constant schema."
)
add_table(
    doc,
    ["Column", "Fuel", "Units"],
    [
        ["demand_fw", "fuelwood", "oven-dry tons"],
        ["demand_ch", "charcoal", "tons of fuelwood-equivalent biomass (kiln yield applied)"],
        ["demand_coal", "coal", "tons"],
        ["demand_gas", "gas", "tons"],
        ["demand_kero", "kerosene", "tons"],
        ["demand_elec", "electric", "MWh"],
        ["demand_biogas", "biogas", "tons"],
        ["demand_eth", "ethanol", "tons"],
        ["demand_pel", "pellets", "tons"],
        ["demand_impfw", "imp_fuelwood", "oven-dry tons"],
        ["demand_impch", "imp_charcoal", "tons of fuelwood-equivalent biomass"],
    ],
)
add_para(
    doc,
    "Note: the 2026–2035 window is fixed in the tool, not derived from today's date. If the "
    "end-year slider is set below 2035 the window is truncated (for example 2026–2030 at an "
    "end year of 2030); the app shows a warning and the actual span is always recorded in "
    "demand_yr_start / demand_yr_end, so read those rather than assuming a full decade.",
    italic=True,
)

doc.add_heading("About the kiln yield in the CSV preamble", level=3)
doc.add_paragraph(
    "The efchratio value in the CSV preamble is the wood-to-charcoal kiln yield set in the "
    "Per-capita rates input tab (§ 3). It applies to charcoal rows of fuel_cons_tons: those "
    "rows are multiplied by the kiln yield so they represent the upstream wood biomass "
    "burned to produce the charcoal — the right unit for forest-impact / biomass-supply "
    "analyses, not stove-side charcoal mass."
)

# ---------- Section 2 ----------

doc.add_heading("Section 2 — Inputs: Fuel Shares", level=1)
doc.add_paragraph(
    "Click \"See / edit input data\" at the top to switch from outputs to the inputs view. "
    "The Inputs view has two sub-tabs; Fuel shares is the first."
)
add_screenshot_placeholder(
    doc,
    "Inputs view with Fuel shares sub-tab open",
    "docs/screenshots/05-inputs-fuel-shares-tab.png",
)

doc.add_heading("2.1 The editable share table", level=2)
doc.add_paragraph(
    "The table shows the fraction of population using each fuel for every "
    "(iso3, country, area, fuel, year) in the current filter selection. The editable "
    "column is population_share (values 0–1); all identifier columns are locked."
)
add_screenshot_placeholder(
    doc,
    "Share editor with a few cells being edited",
    "docs/screenshots/06-share-editor.png",
)
doc.add_paragraph(
    "Five fuels — biogas, ethanol, pellets, imp_fuelwood and imp_charcoal — are absent from "
    "the UN/WHO household survey that supplies the default shares. They are injected with a "
    "population share of 0 for every country, area and year so they can be edited here like "
    "any other fuel: select one in the sidebar, enter shares in this table, and it flows "
    "through consumption and emissions normally. Per-capita rates and emission factors for "
    "them are already present in the reference data, so nothing else needs supplying."
)
doc.add_paragraph("To edit:")
doc.add_paragraph("Click a cell in the population_share column and type a new value, OR paste a column from Excel.", style="List Number")
doc.add_paragraph("Click Save edits to commit. The status banner confirms the row count applied.", style="List Number")
doc.add_paragraph(
    "Edits are persistent for the rest of the session: they flow into the consumption "
    "calculation the moment you save, and the metadata sheet of any subsequent Excel "
    "export records that in-app edits were made."
)
add_para(
    doc,
    "Note: the editor switches to a read-only preview if the current selection produces "
    "more than 8,000 rows. Narrow your sidebar filters (fewer countries or a shorter year "
    "range) to drop below that threshold and re-enable editing.",
    italic=True,
)

doc.add_heading("2.2 Data source ribbon and citations", level=2)
doc.add_paragraph(
    "Below the table is a coloured ribbon showing the current active data source — blue "
    "for the default Stoner et al. (2021) data, yellow when a custom dataset has been "
    "loaded or in-app edits have been applied. Expand \"View Full Citations\" to see the "
    "bibliographic entries for whatever is currently in use, including any "
    "custom-projection citation you provided."
)
add_screenshot_placeholder(
    doc,
    "Source ribbon + expanded citations",
    "docs/screenshots/07-source-ribbon.png",
)

doc.add_heading("2.3 The four action buttons", level=2)
doc.add_paragraph(
    "Below the citations are four buttons; the last two only appear when the "
    "corresponding state is active."
)
add_screenshot_placeholder(
    doc,
    "The four action buttons row",
    "docs/screenshots/08-share-action-buttons.png",
)
add_table(
    doc,
    ["Button", "What it does"],
    [
        ["Upload custom dataset",
         "Opens a modal to upload your own fuel-share CSV/XLSX (replaces the default Stoner et al. data wholesale). The required columns are iso3, country, region, area, fuel, year, population_share. Region is re-derived from iso3 on upload so user-supplied region values are normalised."],
        ["Customize year projections",
         "Opens a modal to set custom values for a specific year (e.g. 2030); the rest of the timeseries is linearly interpolated between the 1990 baseline and the custom year, then to 2050. Provide either the values inline in a data editor or by uploading a CSV. Either way, a data-source citation is required."],
        ["Revert to default WHO data",
         "Only visible after a custom dataset has been loaded. Restores the bundled Stoner et al. data."],
        ["Clear custom projections",
         "Only visible after custom-year projections are active. Removes them and returns to the baseline projections."],
    ],
)

add_para(
    doc,
    "Note on the 0-share fuels: interpolation anchors to the first and last baseline years, "
    "so setting a custom value for a fuel whose baseline is 0 everywhere produces a ramp up "
    "to your custom year and then a decay back to 0 by 2050. To hold such a fuel at a steady "
    "share, edit the values directly in the Fuel shares table (§ 2.1) instead.",
    italic=True,
)

doc.add_heading("Custom upload — minimum requirements", level=3)
doc.add_paragraph("File must be .csv or .xlsx with the header row matching iso3, country, region, area, fuel, year, population_share (case-insensitive; the loader lowercases column names).", style="List Bullet")
doc.add_paragraph("population_share values must be between 0 and 1.", style="List Bullet")
doc.add_paragraph("A free-text data-source description is required at upload time (it shows up in the citations expander and the Metadata sheet of every subsequent Excel export).", style="List Bullet")
doc.add_paragraph("Any Overall rows in the upload are dropped; overall is always derived as urban + rural downstream. Use lowercase urban/rural (case is normalised automatically).", style="List Bullet")
doc.add_paragraph("Fuels your file omits are backfilled automatically: any of biogas, ethanol, pellets, imp_fuelwood or imp_charcoal that is missing is added at a 0 share, so the fuel list stays consistent with the default dataset.", style="List Bullet")

# ---------- Section 3 ----------

doc.add_heading("Section 3 — Inputs: Per-Capita Rates", level=1)
doc.add_paragraph("The second Inputs sub-tab.")
add_screenshot_placeholder(
    doc,
    "Per-capita rates sub-tab",
    "docs/screenshots/09-per-capita-tab.png",
)

doc.add_heading("3.1 The editable rate table", level=2)
doc.add_paragraph(
    "Each row gives the per-capita annual fuel consumption for one (region, fuel) pair. "
    "The editable column is pc_fuel; fuel, region, and pc_fuel_units are locked. The "
    "table is filtered to the fuels selected in the sidebar so you only see the rows "
    "relevant to your current run."
)
add_screenshot_placeholder(
    doc,
    "Per-capita editor showing rows for a couple of regions",
    "docs/screenshots/10-per-capita-editor.png",
)
doc.add_paragraph(
    "Per-capita rates are keyed on (region, fuel) — not on area or year — so a single edit "
    "applies to every (country, area, year) row in that region/fuel combo. This is by "
    "design: per-capita rates are intended as regional structural assumptions, not "
    "country-level observations."
)
doc.add_paragraph(
    "The pc_fuel_units column shows the unit each rate is expressed in (varies by fuel — "
    "MWh/person-year for electric, oven-dry tons/person-year for fuelwood/imp_fuelwood, "
    "tons/person-year for the rest)."
)
doc.add_paragraph(
    "Click Save edits to apply your changes. As with the share editor, edits are "
    "persisted for the rest of the session and noted in the Excel metadata sheet."
)

doc.add_heading("3.2 Wood-to-charcoal kiln yield (efchratio)", level=2)
doc.add_paragraph(
    "Below the per-capita table is the Wood-to-charcoal kiln yield input — a single "
    "numeric field. This is the value exported as efchratio in the CSV preamble (§ 1.4)."
)
add_screenshot_placeholder(
    doc,
    "The kiln yield input + caption block",
    "docs/screenshots/11-kiln-yield.png",
)
add_table(
    doc,
    ["Detail", "Value"],
    [
        ["What it represents",
         "Kilograms of dry wood needed to produce 1 kg of charcoal in a kiln."],
        ["Where it's applied",
         "The fuel_cons_tons column for charcoal rows only. After computing the stove-side charcoal mass, the value is multiplied by the kiln yield to give the upstream wood biomass."],
        ["Default",
         "6 (typical traditional earthen-kiln yield in sub-Saharan Africa — FAO regional reference)."],
        ["Typical range",
         "3–5 for more efficient kilns; 1 if you want to keep charcoal rows as stove-side mass."],
    ],
)
doc.add_paragraph(
    "Why this matters: charcoal kilns waste most of their input wood as heat. A factor of "
    "6 means that producing 1 ton of charcoal consumes ~6 tons of dry wood — the right "
    "number for forest-impact / fNRB analyses. The caption block under the input links to "
    "the UNFCCC fNRB assessment that informed the default."
)

doc.add_heading("3.3 Non-residential wood & charcoal consumption", level=2)
doc.add_paragraph(
    "Below the kiln yield is an optional adjustment for non-residential demand. The "
    "fuel-share data covers household cooking only; this uplift accounts for restaurants, "
    "schools, prisons, bakeries and other institutional or small-commercial users drawing on "
    "the same wood and charcoal supply."
)
add_screenshot_placeholder(
    doc,
    "The non-residential checkbox with the two percentage inputs revealed",
    "docs/screenshots/11b-non-residential.png",
)
add_table(
    doc,
    ["Detail", "Value"],
    [
        ["How to enable",
         "Tick \"Include non-residential wood and charcoal consumption\". Unticked by default, "
         "so outputs are household-only unless you opt in."],
        ["Controls",
         "Two percentage inputs appear when ticked: wood uplift (default 10%) and charcoal "
         "uplift (default 20%)."],
        ["Where it's applied",
         "fuel_cons_tons is multiplied by (1 + uplift/100) — the wood figure for fuelwood and "
         "imp_fuelwood, the charcoal figure for charcoal and imp_charcoal. Applied after the "
         "kiln yield, so charcoal scales on its fuelwood-equivalent basis."],
        ["Effect on other columns",
         "num_fuel_users_thousands is unchanged — this is extra demand on the same fuel "
         "supply, not extra people. Emissions follow the uplifted consumption automatically."],
        ["Where it's recorded",
         "A banner above the consumption table while active; the nonres_wood_pct and "
         "nonres_charcoal_pct columns of the CSV preamble; and the \"Non-Residential "
         "Consumption\" row of both Excel metadata sheets."],
    ],
)
doc.add_paragraph(
    "The defaults of 10% for wood and 20% for charcoal are starting points, not measured "
    "values — the non-residential share varies widely by country and settlement type. Any "
    "alternative entry should be supported by a documented field-based assessment."
)

doc.add_heading("3.4 Custom per-capita upload", level=2)
doc.add_paragraph(
    "If the default per-capita data is missing or you want to replace it, the Upload "
    "custom rates button (visible at the top of the tab when no default is loaded, and "
    "via a modal action when it is) accepts a .csv or .xlsx with columns fuel, region, "
    "pc_fuel, pc_fuel_units. Rows with an empty region value are treated as global and "
    "expanded into every region at load time. As with the share upload, a data-source "
    "citation is required and recorded in the metadata sheet."
)

# ---------- Section 4 ----------

doc.add_heading("Section 4 — Emissions Output", level=1)
doc.add_paragraph(
    "Back in the Outputs view, the second sub-tab is Total emissions. The emissions "
    "calculation uses the consumption results from Section 1 as input: "
    "total_GHG = fuel_cons_tons × em_intens_GHG for each of CO2, CH4, and N2O."
)
add_screenshot_placeholder(
    doc,
    "Emissions tab with the radio toggle visible",
    "docs/screenshots/12-emissions-tab.png",
)

doc.add_heading("4.1 Two views — toggle at the top", level=2)
doc.add_paragraph(
    "A radio control at the top of the tab switches between two views of the same "
    "emissions calculation:"
)
add_table(
    doc,
    ["View", "Row granularity", "Columns"],
    [
        ["Per-fuel detail",
         "(iso3, country, area, fuel, year)",
         "fuel_cons_tons, em_intens_CO2/CH4/N2O, total_CO2/CH4/N2O"],
        ["Country / area summary (CO2-eq)",
         "(iso3, country, area, year) — summed across all fuels",
         "total_CO2, total_CH4, total_N2O, total_CO2eq"],
    ],
)
add_screenshot_placeholder(
    doc,
    "Per-fuel detail table",
    "docs/screenshots/13-emissions-per-fuel.png",
)
add_screenshot_placeholder(
    doc,
    "Country/area CO2-eq summary table",
    "docs/screenshots/14-emissions-summary.png",
)
doc.add_paragraph("The summary view computes:")
add_code(doc, "total_CO2eq = total_CO2 + 28 × total_CH4 + 265 × total_N2O")
doc.add_paragraph(
    "using IPCC AR5 GWP-100 factors — the UNFCCC Enhanced Transparency Framework default "
    "for reporting after 2024."
)

doc.add_heading("4.2 Units and electricity treatment", level=2)
doc.add_paragraph("All total_* columns are in absolute tons of GHG.", style="List Bullet")
doc.add_paragraph(
    "For non-electric fuels, intensities are mass ratios (kg GHG / kg fuel = tons/ton) "
    "sourced from the per-fuel intensity file.",
    style="List Bullet",
)
doc.add_paragraph(
    "For electricity, per-country Combined Margin grid emission factors (from the UNFCCC "
    "IFI TWG methodology list) are used; the loader divides the source gCO2/kWh by 1,000 "
    "so the in-memory value lines up dimensionally with fuel_cons_tons for electric rows "
    "(which is in MWh).",
    style="List Bullet",
)
doc.add_paragraph(
    "CH4 and N2O for electricity are set to 0 because country-level data isn't available; "
    "for most grids these gases contribute less than 5% of CO2-equivalent, so the omission "
    "has limited impact on the summary view but is an important caveat to keep in mind "
    "for grids where biomass cofiring is significant.",
    style="List Bullet",
)
doc.add_paragraph(
    "The on-screen caption below each table restates these units and the calculation explicitly."
)

doc.add_heading("4.3 Missing-intensity diagnostic", level=2)
doc.add_paragraph(
    "If any rows in the per-fuel detail have no matching emissions intensity (likely "
    "cause: a fuel name in your share/per-capita data doesn't match anything in the "
    "intensity files), a yellow \"N rows have no matching emissions intensity\" expander "
    "appears under the table. Open it to see the count grouped by fuel and identify which "
    "fuel name needs reconciling."
)
add_screenshot_placeholder(
    doc,
    "The warning expander with the missing-fuel summary",
    "docs/screenshots/15-emissions-missing-warning.png",
)

doc.add_heading("4.4 Downloading the emissions data", level=2)
doc.add_paragraph(
    "The Download Excel Workbook button at the bottom of the emissions tab produces a "
    "three-sheet workbook:"
)
add_table(
    doc,
    ["Sheet", "Content"],
    [
        ["Metadata & Sources",
         "Run parameters, intensity sources, GWP factors, units note."],
        ["Total Emissions",
         "The per-fuel detail table (regardless of which view is on screen)."],
        ["Summary by Country-Area",
         "The aggregated CO2-eq summary."],
    ],
)
doc.add_paragraph(
    "The filename stem is independent from the consumption export — it starts as "
    "cooking_fuel_emissions_<start_year>_<end_year> but is editable in the text input "
    "above the button."
)
add_screenshot_placeholder(
    doc,
    "Emissions download button + filename input",
    "docs/screenshots/16-emissions-download.png",
)
doc.add_paragraph(
    "There is no CSV bundle for the emissions tab at present; the Excel workbook is the "
    "only download option here."
)

# ---------- Appendix ----------

doc.add_heading("Appendix — Quick reference", level=1)

doc.add_heading("Where each parameter lives", level=2)
add_table(
    doc,
    ["Parameter", "Where to set it", "Where it shows up in exports"],
    [
        ["selected_countries, selected_fuels, selected_areas", "Sidebar",
         "Metadata sheet, CSV preamble (in CSV bundles), filter applied to data rows"],
        ["start_year", "Fixed at 2000 — not user-editable", "Metadata sheet \"Year Range\""],
        ["end_year", "Sidebar slider",
         "Metadata sheet \"Year Range\", end_year in CSV preamble; also bounds the demand window"],
        ["efchratio (kiln yield)", "Per-capita rates tab → Wood-to-charcoal kiln yield",
         "Metadata sheet, efchratio in CSV preamble, applied to charcoal rows of fuel_cons_tons"],
        ["Non-residential uplift (wood %, charcoal %)",
         "Per-capita rates tab → Non-residential wood & charcoal consumption",
         "Metadata sheet \"Non-Residential Consumption\"; nonres_wood_pct and nonres_charcoal_pct in CSV preamble (0 when off)"],
        ["Decadal demand window (2026–2035)",
         "Fixed in the tool; truncated by the end-year slider",
         "Metadata sheet \"Decadal Demand Window\"; demand_yr_start / demand_yr_end in CSV preamble"],
        ["Total demand per fuel", "Computed — not user-set",
         "demand_fw, demand_ch, … (all 11 fuels) in CSV preamble"],
        ["Custom share dataset / projections", "Fuel shares tab → action buttons",
         "Metadata sheet records the source description and citation"],
        ["Per-capita inline edits", "Per-capita rates tab → data editor + Save",
         "Metadata sheet flag \"Per Capita — In-App Edits: Yes\""],
        ["Share inline edits", "Fuel shares tab → data editor + Save",
         "Metadata sheet flag \"Population Share — In-App Edits: Yes\""],
    ],
)

doc.add_heading("Conventional units cheat sheet", level=2)
add_table(
    doc,
    ["Quantity", "Unit"],
    [
        ["num_fuel_users_thousands", "thousands of people"],
        ["per_capita_fuel_cons",
         "varies by fuel (see column pc_fuel_units on the per-capita tab)"],
        ["fuel_cons_tons", "absolute tons"],
        ["total_CO2, total_CH4, total_N2O, total_CO2eq", "absolute tons of GHG"],
        ["efchratio", "kg wood / kg charcoal (dimensionless)"],
        ["population_share", "fraction in [0, 1]"],
        ["demand_<fuel> (CSV preamble)",
         "tons; MWh for demand_elec; fuelwood-equivalent tons for charcoal"],
        ["nonres_wood_pct, nonres_charcoal_pct", "percent (0 when the uplift is off)"],
    ],
)

# ---------- save ----------

out_path = "USER_GUIDE.docx"
doc.save(out_path)
print(f"Wrote {out_path}")
